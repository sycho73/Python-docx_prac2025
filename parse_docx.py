from docx import Document
import os
from zipfile import ZipFile   #docx는 zip파일이라 이걸로 이미지 접근 가능

def extract_images(docx_path, output_dir="images"): #같은 디렉토리에 추출한 이미지 저장 함수
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

        saved_images=[]

        with ZipFile(docx_path, 'r') as docx_zip:
            # docx의 word/media/ 경로에 그림들이 저장되어있음
            image_files = [f for f in docx_zip.namelist() if f.startswith("word/media/")]

            for idx, image_name in enumerate(image_files, 1):
                image_data=docx_zip.read(image_name)
                extension=os.path.splitext(image_name)[1]
                image_path=os.path.join(output_dir,f"image{idx}{extension}")

                with open(image_path, 'wb') as img_file:
                    img_file.write(image_data)
                    saved_images.append(image_path)
        return saved_images

def parse_docx_print(filename="test.docx"):     #텍스트(표 포함) 파싱 후 출력 함수
    document = Document(filename)

    print("\n 텍스트 내용")
    print("-"*50)
    for para in document.paragraphs:
        text=para.text.strip()
        if text:
            print(text)
    print("\n 표 내용")
    print("-"*50)
    for table_idx, table in enumerate(document.tables):
        print(f"[표 {table_idx+1}]")
        for row in table.rows:
            cells=[cell.text.strip() for cell in row.cells]
            print(" | ".join(cells))
        print("-"*30)

def main(docx_filename="test.docx"):  #메인 실행 함수
    if not os.path.exists(docx_filename):
        print(f"파일이 존재하지 않습니다: {docx_filename}")
        return

    parse_docx_print(docx_filename)

    image_paths=extract_images(docx_filename)
    if image_paths:
        for path in image_paths:
            print(f"저장된 이미지: {path}")
    
    else:
        print("문서에 그림이 없습니다")


if __name__=="__main__":
    main()